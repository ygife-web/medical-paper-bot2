"""
Word出力モジュール

python-docxを使用して、論文要約をWord文書として出力する。
見出し、テーブル、参考文献リスト付きの構造化文書を生成。
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

from pubmed_searcher import Paper

logger = logging.getLogger(__name__)


class WordGenerator:
    """Word文書生成クラス"""

    def __init__(self, config: dict):
        """
        初期化

        Args:
            config: config.yamlから読み込んだ設定辞書
        """
        self.config = config
        self.specialty_name = config.get("specialty_name", "医学")
        self.output_config = config.get("output", {})

    def generate(
        self,
        papers: list[Paper],
        output_path: Optional[str] = None
    ) -> str:
        """
        論文要約をWord文書として生成する

        Args:
            papers: 要約済み論文リスト（優先度順）
            output_path: 出力パス（Noneの場合、設定値を使用）

        Returns:
            生成されたファイルパス
        """
        # 出力パスの決定
        if output_path is None:
            output_dir = Path(self.output_config.get("directory", "output"))
            output_dir.mkdir(exist_ok=True)
            date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = self.output_config.get(
                "filename_format", "週刊医学論文レビュー_{date}.docx"
            ).replace("{date}", date_str)
            output_path = str(output_dir / filename)

        doc = Document()
        self._setup_styles(doc)
        self._add_header(doc, papers)
        self._add_summary_index(doc, papers)
        self._add_papers(doc, papers)
        self._add_summary_table(doc, papers)
        self._add_references(doc, papers)

        doc.save(output_path)
        logger.info(f"Word文書を生成しました: {output_path}")
        return output_path

    def _setup_styles(self, doc: Document):
        """文書スタイルを設定する"""
        # デフォルトフォント設定
        style = doc.styles["Normal"]
        font = style.font
        font.name = "游明朝"
        font.size = Pt(10.5)

        # 見出し1スタイル
        h1 = doc.styles["Heading 1"]
        h1.font.name = "游ゴシック"
        h1.font.size = Pt(16)
        h1.font.color.rgb = RGBColor(0, 51, 102)
        h1.font.bold = True

        # 見出し2スタイル
        h2 = doc.styles["Heading 2"]
        h2.font.name = "游ゴシック"
        h2.font.size = Pt(13)
        h2.font.color.rgb = RGBColor(0, 76, 153)
        h2.font.bold = True

        # 見出し3スタイル
        h3 = doc.styles["Heading 3"]
        h3.font.name = "游ゴシック"
        h3.font.size = Pt(11)
        h3.font.color.rgb = RGBColor(0, 102, 153)
        h3.font.bold = True

    def _add_header(self, doc: Document, papers: list[Paper]):
        """文書ヘッダーを追加する"""
        # タイトル
        title = doc.add_heading("週刊 医学論文レビュー", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # サブタイトル
        now = datetime.now()
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            f"{self.specialty_name} 新着論文サマリー\n"
            f"作成日: {now.strftime('%Y年%m月%d日')}（{now.strftime('%A')}）"
        )
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)

        # 区切り線
        doc.add_paragraph("─" * 50)

        # 概要
        overview = doc.add_paragraph()
        overview.add_run(
            f"今週の注目論文トップ{len(papers)}を選出しました。\n"
            f"冒頭にサマリーインデックスを掲載し、全{len(papers)}本を詳細に解説します。"
        ).font.size = Pt(10)

    def _add_papers(self, doc: Document, papers: list[Paper]):
        """各論文セクションを追加する"""
        detailed_n = self.config.get("search", {}).get("detailed_top_n", 3)

        for i, paper in enumerate(papers):
            is_detailed = (i < detailed_n)

            # AI要約から重要度を抽出
            index_info = self._extract_index_info(paper)
            importance = index_info.get("importance", "")

            # セクション見出し
            doc.add_heading(
                f"#{paper.priority_rank} 【{importance}】 {paper.title}",
                level=1
            )

            # 論文基本情報テーブル
            self._add_paper_info_table(doc, paper)

            # 選出理由
            if hasattr(paper, '_selection_reason') and paper._selection_reason:
                reason_para = doc.add_paragraph()
                reason_run = reason_para.add_run(paper._selection_reason)
                reason_run.font.size = Pt(9)
                reason_run.font.italic = True
                reason_run.font.color.rgb = RGBColor(100, 100, 100)

            # AI要約
            if paper.summary and paper.summary.get("content"):
                content = paper.summary["content"]
                self._add_markdown_content(doc, content)
            else:
                doc.add_paragraph("要約は生成されませんでした。")

            # 区切り
            if i < len(papers) - 1:
                doc.add_paragraph("─" * 50)

    def _add_paper_info_table(self, doc: Document, paper: Paper):
        """論文基本情報テーブルを追加する"""
        table = doc.add_table(rows=5, cols=2)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 著者表示（最大5名）
        if len(paper.authors) > 5:
            author_str = ", ".join(paper.authors[:5]) + " et al."
        else:
            author_str = ", ".join(paper.authors)

        # 論文タイプ
        pub_type_str = ", ".join(paper.pub_types) if paper.pub_types else "N/A"

        # テーブルデータ
        data = [
            ("著者", author_str),
            ("ジャーナル", paper.journal),
            ("出版日", paper.pub_date),
            ("論文タイプ", pub_type_str),
            ("DOI", paper.doi if paper.doi else "N/A"),
        ]

        for row_idx, (label, value) in enumerate(data):
            # ラベルセル
            cell_label = table.cell(row_idx, 0)
            cell_label.text = label
            for paragraph in cell_label.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

            # 値セル
            cell_value = table.cell(row_idx, 1)
            cell_value.text = value
            for paragraph in cell_value.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

        # テーブル幅設定
        for row in table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(13)

        doc.add_paragraph()  # スペース

    def _add_markdown_content(self, doc: Document, content: str):
        """
        マークダウン形式のAI要約をWord文書に追加する

        見出し（##）、リスト（-、*）、太字（**）を変換
        """
        lines = content.split("\n")

        for line in lines:
            stripped = line.strip()

            if not stripped:
                continue

            # 見出し2（##）
            if stripped.startswith("## "):
                heading_text = stripped[3:].strip()
                doc.add_heading(heading_text, level=2)

            # 見出し3（###）
            elif stripped.startswith("### "):
                heading_text = stripped[4:].strip()
                doc.add_heading(heading_text, level=3)

            # リスト項目
            elif stripped.startswith("- ") or stripped.startswith("* "):
                item_text = stripped[2:].strip()
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(para, item_text)

            # 通常テキスト
            else:
                para = doc.add_paragraph()
                self._add_formatted_text(para, stripped)

    def _add_formatted_text(self, paragraph, text: str):
        """
        太字（**text**）を含むテキストをWord段落に追加する
        """
        # **text** パターンを分割
        parts = re.split(r'(\*\*.*?\*\*)', text)

        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                # 太字テキスト
                run = paragraph.add_run(part[2:-2])
                run.font.bold = True
                run.font.size = Pt(10.5)
            else:
                # 通常テキスト
                if part:
                    run = paragraph.add_run(part)
                    run.font.size = Pt(10.5)

    def _add_summary_table(self, doc: Document, papers: list[Paper]):
        """末尾のサマリーテーブルを追加する"""
        doc.add_heading("今週の論文一覧", level=1)

        # ヘッダー行 + データ行
        table = doc.add_table(rows=1 + len(papers), cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # ヘッダー
        headers = ["優先度", "論文", "ジャーナル / デザイン", "一言要約", "実臨床への影響"]
        for col_idx, header in enumerate(headers):
            cell = table.cell(0, col_idx)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)

        # データ行
        for row_idx, paper in enumerate(papers, 1):
            # 優先度
            if paper.priority_rank <= 3:
                priority = f"★ #{paper.priority_rank}"
            else:
                priority = f"#{paper.priority_rank}"
            table.cell(row_idx, 0).text = priority

            # 論文タイトル（短縮）
            title = paper.title[:60] + "..." if len(paper.title) > 60 else paper.title
            table.cell(row_idx, 1).text = title

            # ジャーナル / デザイン
            pub_type = paper.pub_types[0] if paper.pub_types else "N/A"
            table.cell(row_idx, 2).text = f"{paper.journal}\n{pub_type}"

            # 一言要約（AI要約の最初のセクション）
            one_liner = self._extract_one_liner(paper)
            table.cell(row_idx, 3).text = one_liner

            # 実臨床への影響
            impact = self._extract_clinical_impact(paper)
            table.cell(row_idx, 4).text = impact

            # フォントサイズ設定
            for col_idx in range(5):
                for paragraph in table.cell(row_idx, col_idx).paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

        # テーブル幅設定
        widths = [Cm(1.5), Cm(4), Cm(3), Cm(4.5), Cm(4)]
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = width

    def _extract_one_liner(self, paper: Paper) -> str:
        """AI要約から一言要約を抽出する"""
        content = paper.summary.get("content", "")
        lines = content.split("\n")

        capture = False
        result = []
        for line in lines:
            if "まず一言で" in line:
                capture = True
                continue
            if capture:
                stripped = line.strip()
                if stripped.startswith("##"):
                    break
                if stripped:
                    result.append(stripped)

        if result:
            return " ".join(result)[:100]
            
        # フォールバック: サマリーインデックスの「結論」を使用
        info = self._extract_index_info(paper)
        return info.get("conclusion", "要約なし")

    def _extract_clinical_impact(self, paper: Paper) -> str:
        """AI要約から臨床的影響を抽出する"""
        content = paper.summary.get("content", "")
        lines = content.split("\n")

        capture = False
        result = []
        for line in lines:
            if "実践メモ" in line or "臨床的に重要" in line:
                capture = True
                continue
            if capture:
                stripped = line.strip()
                if stripped.startswith("##"):
                    break
                if stripped.startswith("- ") or stripped.startswith("* "):
                    result.append(stripped[2:])

        if result:
            # 最初の2項目を返す
            return "; ".join(result[:2])[:100]
            
        # フォールバック: サマリーインデックスの「実用」を使用
        info = self._extract_index_info(paper)
        return info.get("practical", "要約参照")

    def _add_summary_index(self, doc: Document, papers: list[Paper]):
        """冒頭のサマリーインデックスを追加する"""
        doc.add_heading("サマリーインデックス", level=1)
        
        for i, paper in enumerate(papers, 1):
            info = self._extract_index_info(paper)
            importance = info.get("importance", "★★★☆☆")
            conclusion = info.get("conclusion", "要約参照")
            practical = info.get("practical", "要約参照")

            # [番号]. 【★重要度】タイトル
            p_title = doc.add_paragraph()
            p_title.paragraph_format.space_before = Pt(10)
            run_title = p_title.add_run(f"{i}. 【{importance}】 {paper.title}")
            run_title.font.bold = True
            run_title.font.size = Pt(11)

            # ・結論：...
            p_conc = doc.add_paragraph(style="List Bullet")
            p_conc.paragraph_format.left_indent = Cm(0.5)
            self._add_formatted_text(p_conc, f"**結論**：{conclusion}")

            # ・実用：...
            p_prac = doc.add_paragraph(style="List Bullet")
            p_prac.paragraph_format.left_indent = Cm(0.5)
            self._add_formatted_text(p_prac, f"**実用**：{practical}")

        doc.add_page_break()

    def _extract_index_info(self, paper: Paper) -> dict:
        """AI要約からインデックス用情報を抽出する"""
        content = paper.summary.get("content", "")
        info = {
            "importance": "★★★☆☆",
            "conclusion": "要約参照",
            "practical": "要約参照"
        }
        
        # サマリーインデックス情報セクションを探す
        # セクション見出しの揺らぎ（## や ###、前後のスペース等）を許容
        section_match = re.search(r"(?:##+|[*]{2})\s*サマリーインデックス情報.*?\n(.*?)(?=\n#|$)", content, re.DOTALL | re.IGNORECASE)
        
        section_content = section_match.group(1) if section_match else content
            
        # 重要度抽出
        imp_match = re.search(r"重要度.*?([★☆]+)", section_content)
        if imp_match:
            info["importance"] = imp_match.group(1).strip()
                
        # 結論抽出: 結論から実用（またはセクション末尾）までを丸ごと取得
        conc_match = re.search(r"結論(.*?)(?:実用|$)", section_content, re.DOTALL)
        if conc_match:
            val = conc_match.group(1)
            val = re.sub(r"^[:：\s\*・-]*", "", val)
            # 末尾の改行、スペース、装飾、次のリストマーカー（例: \n3. ）等を除去
            val = re.sub(r"[\s\*・\-\d\.]*$", "", val)
            if val:
                info["conclusion"] = val
                
        # 実用抽出: 実用からセクション末尾までを丸ごと取得
        prac_match = re.search(r"実用(.*?)$", section_content, re.DOTALL)
        if prac_match:
            val = prac_match.group(1)
            val = re.sub(r"^[:：\s\*・-]*", "", val)
            val = re.sub(r"[\s\*・\-\d\.]*$", "", val)
            if val:
                info["practical"] = val
                
        return info

    def _add_references(self, doc: Document, papers: list[Paper]):
        """参考文献リストを追加する"""
        doc.add_heading("参考文献", level=1)

        for i, paper in enumerate(papers, 1):
            # 著者表示
            if len(paper.authors) > 3:
                author_str = ", ".join(paper.authors[:3]) + " et al."
            else:
                author_str = ", ".join(paper.authors)

            # 参照フォーマット
            ref_text = (
                f"{i}. {author_str}. {paper.title}. "
                f"{paper.journal}. {paper.pub_date}."
            )

            if paper.doi:
                ref_text += f" doi: {paper.doi}"

            ref_text += f" PMID: {paper.pmid}"

            para = doc.add_paragraph()
            run = para.add_run(ref_text)
            run.font.size = Pt(9)
